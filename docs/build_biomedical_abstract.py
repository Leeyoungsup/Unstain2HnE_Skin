from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
ASSET_DIR = HERE / "assets" / "unstain2hne_wsi_report"
OUTPUT = HERE / "Unstain2HnE_의공학_확장초록.docx"
FIGURE_1 = ASSET_DIR / "biomedical_abstract_figure1.png"

FONT = "Noto Sans CJK KR"
BODY_SIZE = 9.0


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=2,
                  line=1.0, keep=False):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep
    fmt.widow_control = True
    return paragraph


def add_text(container, text, *, size=BODY_SIZE, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=2, line=1.0,
             keep=False, color=None):
    paragraph = container.add_paragraph()
    set_paragraph(
        paragraph, align=align, before=before, after=after, line=line, keep=keep
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_heading(container, number, title):
    paragraph = add_text(
        container, f"{number}. {title}", size=9.2, bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=2, keep=True,
    )
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "4")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "7B1FA2")
    bottom.append(border)
    paragraph._p.get_or_add_pPr().append(bottom)
    return paragraph


def set_cell_margins(cell, top=60, start=80, bottom=40, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def shade_cell(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=7.5, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, align=align, after=0, line=1.0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_result_table(container):
    rows = [
        ("지표", "Baseline", "Virtual H&E"),
        ("Full RGB SSIM ↑", "0.518", "0.527"),
        ("Tissue SSIM ↑", "0.118", "0.149"),
        ("Coarse SSIM ↑", "0.537", "0.582"),
        ("Full PSNR ↑", "13.67 dB", "18.01 dB"),
        ("Tissue ΔE2000 ↓", "37.54", "15.01"),
        ("LPIPS ↓", "0.358", "0.223"),
        ("H spatial corr. ↑", "0.104", "0.527"),
        ("Tissue Dice ↑", "0.864", "0.963"),
    ]
    table = container.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.3), Cm(2.2), Cm(2.6)]
    for row_index, row in enumerate(rows):
        for column, value in enumerate(row):
            cell = table.cell(row_index, column)
            cell.width = widths[column]
            set_cell_margins(cell, 25, 35, 25, 35)
            if row_index == 0:
                shade_cell(cell, "6A1B9A")
                set_cell_text(cell, value, bold=True, size=7.2, color=(255, 255, 255))
            else:
                if row_index % 2 == 0:
                    shade_cell(cell, "F3E5F5")
                set_cell_text(
                    cell, value, bold=(column == 2),
                    align=WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                    size=7.1,
                )
    add_text(
        container,
        "표 2. 2.0 MPP held-out case-level 성능. Baseline은 경쟁 모델이 아니라 입력 Unstain RGB이다.",
        size=6.9, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=3,
    )


def add_registration_table(container):
    rows = [
        ("Slide", "IoU 전", "IoU 후", "ECC", "QC"),
        ("MD22-04155(A1-2)", "0.718", "0.946", "0.992", "Pass"),
        ("MD22-04167(B1-2)", "0.849", "0.972", "0.989", "Pass"),
        ("MD22-04169(B1-2)", "0.849", "0.970", "0.994", "Pass"),
        ("MD22-4228(A1-3)", "0.904", "0.953", "0.985", "Pass"),
        ("WC21-01351(A1-2)", "0.919", "0.930", "0.962", "Pass"),
        ("WC21-01356(B1-1)", "0.933", "0.939", "0.978", "Pass"),
        ("WC21-01356(B1-2)", "0.863", "0.868", "0.937", "Fail"),
        ("WC21-01958(A1-1)", "0.598", "0.938", "0.984", "Pass"),
        ("평균", "0.829", "0.939", "0.978", "7/8"),
    ]
    table = container.add_table(rows=len(rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.8), Cm(1.25), Cm(1.25), Cm(1.25), Cm(1.1)]
    for row_index, row in enumerate(rows):
        for column, value in enumerate(row):
            cell = table.cell(row_index, column)
            cell.width = widths[column]
            set_cell_margins(cell, 28, 25, 28, 25)
            if row_index == 0:
                shade_cell(cell, "37474F")
                set_cell_text(cell, value, bold=True, size=6.8, color=(255, 255, 255))
            else:
                if row_index % 2 == 0:
                    shade_cell(cell, "ECEFF1")
                if value == "Fail":
                    shade_cell(cell, "FFCDD2")
                set_cell_text(
                    cell, value, bold=(row_index == len(rows) - 1 or column == 2),
                    align=WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                    size=6.7,
                )
    add_text(
        container,
        "표 1. Held-out WSI별 affine registration 성능. QC 기준: mask IoU after ≥0.90.",
        size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=3,
    )


def add_picture(container, path, width_cm, caption):
    paragraph = container.add_paragraph()
    set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=1, keep=True)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    add_text(
        container, caption, size=6.8, align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0, after=3, line=1.0,
    )


def add_bullet(container, text):
    paragraph = container.add_paragraph(style=None)
    set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=1, line=1.0)
    paragraph.paragraph_format.left_indent = Mm(3)
    paragraph.paragraph_format.first_line_indent = Mm(-2.4)
    run = paragraph.add_run("• " + text)
    set_run_font(run, size=BODY_SIZE)


def compose_stage_figure():
    stage1 = Image.open(ASSET_DIR / "stage1_structure_epoch0805.png").convert("RGB")
    stage2 = Image.open(ASSET_DIR / "stage2_color_epoch0952.png").convert("RGB")
    target_width = 1800
    images = []
    for image in (stage1, stage2):
        height = round(image.height * target_width / image.width)
        images.append(image.resize((target_width, height), Image.Resampling.LANCZOS))
    gap = 25
    canvas = Image.new("RGB", (target_width, sum(im.height for im in images) + gap), "white")
    y = 0
    for image in images:
        canvas.paste(image, (0, y))
        y += image.height + gap
    canvas.save(FIGURE_1, quality=95)


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— ")
    set_run_font(run, size=7)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(" —")
    set_run_font(run, size=7)


def build_document():
    compose_stage_figure()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(6)
    add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(BODY_SIZE)

    title = add_text(
        doc,
        "염색 전 피부조직 영상의 구조–색상 분리형 2단계 가상 H&E 생성 및 전슬라이드 영상 재구성",
        size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=1.08,
    )
    title.paragraph_format.keep_with_next = True
    add_text(
        doc,
        "Structure–Color Decoupled Two-stage Virtual H&E Generation and Whole-slide Reconstruction from Unstained Skin Tissue",
        size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3, line=1.0, keep=True,
    )
    add_text(
        doc, "이영섭¹ · [공동저자]² · [교신저자]¹*",
        size=9.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1, keep=True,
    )
    add_text(
        doc, "¹[소속 및 학과], ²[공동연구기관]  |  *Corresponding: [e-mail]",
        size=7.8, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, keep=True,
    )

    abstract_table = doc.add_table(rows=1, cols=1)
    abstract_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = abstract_table.cell(0, 0)
    set_cell_margins(cell, 80, 120, 75, 120)
    shade_cell(cell, "F7F3F8")
    add_text(cell, "초 록", size=9.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, keep=True)
    add_text(
        cell,
        "염색 전 조직 영상 기반 가상 H&E는 신속한 조직 확인 가능성을 제공하지만, 입력에서 관찰되지 않는 세포 구조의 생성과 패치 기반 전슬라이드 영상(WSI)의 색상 불연속이 문제이다. 본 연구에서는 구조 생성과 색상 생성을 분리한 2단계 가상염색 방법을 제안하였다. 1단계는 전역 보정된 1채널 optical density(OD)를 입력받는 CycleGAN으로 H&E 구조 OD를 예측하고, 2단계는 예측 OD만으로 RGB H&E를 생성한다. 0.5 MPP의 2048×2048 패치를 동일 시야의 2.0 MPP, 512×512 입력으로 변환하여 학습하였다. 실제 H&E를 moving, Unstain을 fixed 영상으로 정의하고 조직 마스크 기반 affine ECC 정합을 수행했으며, thumbnail 변환행렬을 level-0 및 2.0 MPP 좌표계로 변환하여 정량평가에 적용하였다. WSI 추론에는 50% overlap과 cosine blending을 적용하고, 배경 마스크를 합성 후 전역적으로 처리하였다. 7명 8개 held-out WSI의 129개 조직 타일에서 정합된 실제 H&E와 비교한 결과, virtual H&E는 입력 Unstain baseline 대비 tissue SSIM 0.118에서 0.149, full PSNR 13.67 dB에서 18.01 dB, LPIPS 0.358에서 0.223, ΔE2000 37.54에서 15.01로 개선되었다. Hematoxylin 공간 상관계수는 0.104에서 0.527로 증가하였다. 평균 tissue-mask IoU는 정합 전 0.829에서 정합 후 0.939로 증가하였다. 또한 seam ratio는 1.158에서 1.033으로 감소하여 타일 경계 초과 불연속이 약 79% 줄었다. 제안 방법은 2.0 MPP 조직학적 표현과 WSI 연속성을 개선했으나, 입력에서 누락된 개별 세포의 진위성을 보장하지 않으므로 외부 코호트 및 병리의사 평가가 추가로 필요하다.",
        size=8.0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=2, line=1.0,
    )
    add_text(
        cell,
        "핵심어: Virtual staining, H&E, CycleGAN, optical density, whole-slide image, digital pathology",
        size=7.6, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=0,
    )
    remove_table_borders(abstract_table)

    page1 = doc.add_table(rows=1, cols=2)
    page1.alignment = WD_TABLE_ALIGNMENT.CENTER
    page1.autofit = False
    remove_table_borders(page1)
    left, right = page1.rows[0].cells
    for current in (left, right):
        current.width = Cm(8.9)
        current.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(current, 40, 70, 20, 70)
        current.text = ""

    add_heading(left, "1", "서론")
    add_text(
        left,
        "가상염색은 염색 전 조직에서 H&E와 유사한 영상을 생성하여 염색 시간과 조직 소모를 줄일 가능성이 있다. 그러나 피부조직의 Unstained 영상은 H&E에서 강조되는 핵과 일부 세포 구조를 충분히 포함하지 않으며, 정합된 쌍에도 절편 변형과 국소 오차가 존재한다. 강한 pixel-wise paired loss는 이러한 오차를 학습할 수 있고, 단일 단계 RGB 생성은 구조와 색상 문제를 동시에 해결해야 한다.",
    )
    add_text(
        left,
        "본 연구의 목표는 (1) 구조와 색상 생성을 분리하고, (2) 잔여 정합 오차에 강한 cycle consistency를 유지하며, (3) WSI 타일 경계와 배경 불연속을 최소화하는 것이다.",
    )

    add_heading(left, "2", "재료 및 방법")
    add_text(
        left,
        "Paired Unstained–H&E 피부조직 패치를 case 단위로 train/validation에 분리하였다(seed=42). 원본 0.5 MPP, 2048 px 패치는 동일 물리 시야를 유지하여 2.0 MPP, 512 px로 변환하였다. OD는 영상별 min–max가 아니라 256개 학습 영상으로 보정한 전역 범위를 사용했고, 네트워크 입력은 [−1, 1]로 변환하였다.",
    )
    add_text(
        left,
        "Stage 1은 6개 residual block을 갖는 1채널 ResNet CycleGAN이다. 양방향 adversarial 및 cycle/identity loss를 유지하면서 paired L1, SSIM, gradient, background loss를 결합하였다. 이는 구조 정합 정보를 사용하되 작은 registration error에 대한 과의존을 줄이기 위한 구성이다.",
    )
    add_text(
        left,
        "Stage 2는 predicted H&E OD 한 채널만 입력받아 RGB H&E를 생성한다. RGB, full-resolution SSIM, gradient, Laplacian, adversarial 및 background loss를 사용하고, bilinear upsampling 뒤 residual detail refinement를 적용하였다. 원본 Unstain 또는 real H&E OD teacher 경로는 사용하지 않았으며 blur 기반 loss도 제외하였다.",
    )
    add_bullet(left, "Stage 1 latest checkpoint: epoch 805")
    add_bullet(left, "Stage 2 latest checkpoint: epoch 952")
    add_bullet(left, "평가 해상도: 2.0 MPP")

    add_heading(right, "2.1", "WSI 생성 및 평가 개요")
    add_text(
        right,
        "WSI는 512 px 타일, 256 px overlap(stride 256)으로 추론하였다. 겹침 구간은 flat-centre cosine weight로 합성하고, source-derived background mask는 합성 완료 후 WSI 좌표계에서 한 번만 적용하였다. 조직 픽셀에는 blur를 적용하지 않았으며 pyramidal TIFF는 JPEG quality 95로 저장하였다.",
    )
    add_text(
        right,
        "Held-out 7 cases, 8 slides에서 affine registration된 실제 H&E를 기준으로 129개 조직 타일을 평가하였다. SSIM, PSNR, MAE, ΔE2000, LPIPS, stain concentration, gradient correlation, tissue Dice, FID/KID를 계산했다. case-level bootstrap 95% 신뢰구간과 paired Wilcoxon 검정 및 Benjamini–Hochberg 보정을 사용하였다. Registration tissue-mask IoU 0.90을 QC 기준으로 설정하였다.",
    )
    add_picture(
        right, FIGURE_1, 8.3,
        "그림 1. Stage 1의 OD 구조 변환(상)과 Stage 2의 predicted OD 기반 RGB H&E 생성(하).",
    )

    doc.add_page_break()

    registration_page = doc.add_table(rows=1, cols=2)
    registration_page.alignment = WD_TABLE_ALIGNMENT.CENTER
    registration_page.autofit = False
    remove_table_borders(registration_page)
    left, right = registration_page.rows[0].cells
    for current in (left, right):
        current.width = Cm(8.9)
        current.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(current, 20, 70, 20, 70)
        current.text = ""

    add_heading(left, "3", "Unstain–H&E WSI 정합")
    add_text(
        left,
        "동일 검체의 Unstain WSI를 fixed image, 염색 후 H&E WSI를 moving image로 정의하였다. 염색·세척·재스캔 과정에서 발생하는 평행이동, 회전, 배율 변화 및 전단을 보정하기 위해 mask 기반 affine registration을 사용하였다. 정합은 모델 성능평가 이전에 독립적으로 수행했으며, 생성 Virtual H&E에는 변환을 적용하지 않고 real H&E만 Unstain 좌표계로 이동시켰다.",
    )

    add_heading(left, "3.1", "Thumbnail 및 조직 마스크")
    add_text(
        left,
        "각 WSI를 OpenSlide로 읽고 장축 최대 1,000 px의 RGB thumbnail을 생성하였다. 두 영상의 크기가 다르면 흰색 배경으로 중앙 정렬 padding하여 동일 canvas를 만들었다. RGB를 grayscale로 변환하고 5×5 Gaussian smoothing 후 inverse Otsu thresholding으로 tissue mask를 생성하였다.",
    )
    add_text(
        left,
        "Unstain과 H&E의 명암 분포 차이를 보완하기 위해 fixed mask에는 Otsu threshold offset {0, 2, 4, 6, 8, 10}을 각각 적용하고, moving tissue area와 로그 면적비가 가장 가까운 후보를 선택하였다. 5×5 morphological closing 2회와 opening 1회를 적용하고, 전체 canvas의 0.05% 미만인 연결성분(최소 16 px)은 제거하였다.",
    )

    add_heading(left, "3.2", "Affine ECC 최적화")
    add_text(
        left,
        "이진 tissue mask를 σ=3으로 완만하게 만든 뒤, 두 마스크의 centroid 차이로 x–y translation을 초기화하였다. OpenCV Enhanced Correlation Coefficient(ECC)를 affine motion model로 최적화하였다. 종료 조건은 최대 500 iteration 또는 ΔECC<10⁻⁷이며, ECC 내부 Gaussian filter size는 5를 사용하였다. 최종 2×3 inverse warp는 역행렬을 취해 moving H&E→fixed Unstain 방향의 3×3 homogeneous transform으로 저장하였다.",
    )
    add_text(
        left,
        "Affine model은 translation뿐 아니라 rotation, anisotropic scale 및 shear를 포함한다. Held-out 영상에서 회전 범위는 −3.00°–0.88°, scale-x는 1.000–1.013, scale-y는 0.999–1.018이었다.",
    )

    add_heading(left, "3.3", "Level-0 및 목표 MPP 좌표변환")
    add_text(
        left,
        "Thumbnail 정합행렬에는 각 WSI의 원본 크기와 center-padding offset을 역으로 반영하여 level-0 full-resolution H&E→Unstain 행렬을 계산하였다. 평가 시에는 fixed/moving WSI의 실제 MPP를 이용하여 다음과 같이 2.0 MPP raster 좌표계로 변환하였다.",
    )
    add_text(
        left,
        "M₂.₀ = S_fixed · M_level-0 · S_moving⁻¹",
        size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3,
    )
    add_text(
        left,
        "Real H&E RGB에는 bilinear interpolation을 적용하고, 유효영역 mask에는 nearest-neighbor interpolation을 적용하였다. 이동 후 canvas 밖 영역은 흰색으로 채우고, valid mask 밖의 픽셀은 모든 paired metric에서 제외하였다. 행렬, 원본 크기, MPP, affine parameter와 품질지표는 slide별 JSON으로 저장하였다.",
    )

    add_heading(right, "3.4", "정합 품질관리")
    add_picture(
        right,
        ASSET_DIR / "registration_WC21-01958_A1-1.png",
        8.4,
        "그림 2. 대표 WSI 정합. Moving H&E, fixed Unstain, 정합 전 중첩, aligned H&E, mask overlap 및 정합 후 중첩을 순서대로 표시하였다.",
    )
    add_registration_table(right)
    add_text(
        right,
        "Held-out 8개 WSI의 평균 mask IoU는 0.829±0.115에서 0.939±0.033으로 증가했고, 평균 ECC score는 0.978±0.019였다. WC21-01958(A1-1)은 IoU가 0.598에서 0.938로 가장 크게 개선되었다. WC21-01356(B1-2)는 정합 후 IoU가 0.868로 QC 기준 0.90을 충족하지 못했다.",
    )
    add_text(
        right,
        "평가 결과는 (1) 전체 8개 WSI 분석과 (2) registration QC-pass 7개 WSI sensitivity analysis로 분리해 저장하였다. ECC가 높더라도 mask IoU가 낮으면 구조적으로 잘못 정렬될 수 있으므로 두 지표를 함께 확인하였다. Affine registration은 전체 조직 형태의 차이를 보정하지만 절편 과정에서 발생한 국소 비선형 변형이나 실제 세포 소실은 복원하지 못한다.",
    )
    add_text(
        right,
        "정합 파이프라인: H&E/Unstain thumbnail → center padding → stain-independent tissue mask → centroid translation initialization → affine ECC → thumbnail-to-level-0 matrix conversion → target-MPP warp → valid-mask 및 IoU/ECC QC",
        size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=3,
    )

    doc.add_page_break()

    page2 = doc.add_table(rows=1, cols=2)
    page2.alignment = WD_TABLE_ALIGNMENT.CENTER
    page2.autofit = False
    remove_table_borders(page2)
    left, right = page2.rows[0].cells
    for current in (left, right):
        current.width = Cm(8.9)
        current.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(current, 20, 70, 20, 70)
        current.text = ""

    add_heading(left, "4", "가상염색 결과")
    add_text(
        left,
        "8개 WSI 모두 추론에 성공하였다. Virtual H&E는 baseline보다 색상·지각적 유사도와 hematoxylin 공간 분포가 개선되었다. Tissue SSIM, coarse SSIM, PSNR, MAE, ΔE2000, LPIPS, H spatial correlation과 tissue Dice의 case-paired 개선은 FDR 보정 후 q≈0.022였다. FID는 232.71에서 106.80, KID는 0.1393에서 0.0357로 감소하였다.",
    )
    add_result_table(left)
    add_picture(
        left,
        ASSET_DIR / "slide_MD22-04155_A1-2.png",
        8.3,
        "그림 3. 대표 WSI의 Unstain 입력, virtual H&E, 정합 real H&E 및 조직 영역 오차.",
    )
    add_text(
        left,
        "Registration QC는 8개 중 7개가 통과했으며, WC21-01356(B1-2)는 mask IoU 0.868로 sensitivity analysis에서 제외 대상으로 표시하였다. 따라서 전체 결과와 QC-pass 결과를 함께 보고해야 한다.",
        size=7.8,
    )

    add_picture(
        right,
        ASSET_DIR / "cohort_metric_summary.png",
        7.9,
        "그림 4. Held-out case-level 평균과 bootstrap 95% 신뢰구간.",
    )
    add_picture(
        right,
        ASSET_DIR / "seam_audit.png",
        8.2,
        "그림 5. 50% overlap cosine blending 적용 전후의 WSI seam 분석.",
    )

    add_heading(right, "5", "고찰")
    add_text(
        right,
        "Stage 1/2 분리는 구조적 OD 생성과 H&E 색상 생성을 서로 다른 목적함수로 최적화할 수 있게 했다. 50% overlap은 추론량을 약 3배 증가시키지만 seam ratio를 1.158에서 1.033으로 낮춰 타일 경계 초과 불연속을 약 79% 감소시켰다. Full grayscale SSIM은 0.502에서 0.495로 소폭 감소했고 Laplacian energy error도 악화되어, 세포 수준 고주파 표현에는 추가 개선이 필요하다.",
    )
    add_text(
        right,
        "Tissue SSIM은 0.149로 full SSIM보다 낮았다. 이는 tissue-only metric이 잔여 registration error와 절편 간 구조 차이에 민감하기 때문이다. 반대로 배경이 포함된 SSIM만 사용하면 성능이 과대평가될 수 있다. 또한 생성된 세포 형태가 plausible하더라도 입력에 없는 개별 세포의 실제 존재를 증명하지는 않는다.",
    )

    add_heading(right, "6", "결론")
    add_text(
        right,
        "제안한 구조–색상 분리형 2단계 모델은 2.0 MPP에서 Unstained 피부조직을 조직학적 virtual H&E로 변환하고 WSI 연속성을 개선하였다. 향후 외부 코호트, 고해상도 핵 검출 일치도, 병리의사 blinded reader study 및 downstream 조직분할·병변분류 검증이 필요하다.",
    )

    add_heading(right, "참고문헌", "")
    references = [
        "[1] Zhu JY, et al. Unpaired Image-to-Image Translation Using Cycle-Consistent Adversarial Networks. ICCV, 2017.",
        "[2] Wang Z, et al. Image Quality Assessment: From Error Visibility to Structural Similarity. IEEE TIP, 2004.",
        "[3] Zhang R, et al. The Unreasonable Effectiveness of Deep Features as a Perceptual Metric. CVPR, 2018.",
        "[4] Heusel M, et al. GANs Trained by a Two Time-Scale Update Rule Converge to a Local Nash Equilibrium. NeurIPS, 2017.",
    ]
    for reference in references:
        add_text(right, reference, size=6.7, align=WD_ALIGN_PARAGRAPH.LEFT, after=1)

    core = doc.core_properties
    core.title = "Unstain2HnE Two-stage Virtual H&E — Biomedical Engineering Abstract"
    core.subject = "Expanded biomedical engineering abstract"
    core.author = "이영섭 외"
    core.keywords = "virtual staining, H&E, CycleGAN, WSI, digital pathology"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
